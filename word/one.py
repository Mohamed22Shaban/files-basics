from audioop import add
from turtle import width
import docx
from pathlib import Path 


read_file = docx.Document(Path.home() / Path('D:','word','academy_1.docx'))

# how many paragraph
print(len(read_file.paragraphs))

# to print the first paragraph
print(read_file.paragraphs[0].text)

# to obtain paragraph's division 
print(len(read_file.paragraphs[2].runs))
print(read_file.paragraphs[0].runs[1].text)

# print the whole text
def get_text(filename):
    read_file = docx.Document(filename)
    l = [] 
    for para in read_file.paragraphs:
        l.append(para.text)
    return '\n'.join(l)

print(get_text(Path.home() / Path('D:','word','academy_1.docx')))

#===========================================================================

# Write To Word

#1) create object
mydoc = docx.Document()
mydoc.add_paragraph('this is the first paragraph i added')
mydoc.add_paragraph('the second paragraph is full')

# to write by arabic shall call "WD_PARAGRAPH_ALIGNMENT" =RIGHT OR LEFT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
arabic = mydoc.add_paragraph('اول نص هيكون عن الرحمه')
arabic.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

# to add sentence to the clause 
third_paragraph = mydoc.add_paragraph('this is the third paragraph ')
third_paragraph.add_run('this is section in the third paragraph')

# add heading

mydoc.add_heading('Number One', 0) # => size of heading start from 0 to 6
mydoc.add_heading('Number One', 2) 
mydoc.add_heading('Number One', 4)



# style 
print(mydoc.paragraphs[1].style)   # will give normal because it doesnot has any formatting
mydoc.paragraphs[1].style = mydoc.styles['Heading 1']

#could delete style
mydoc.paragraphs[1].style.delete()


# add image
mydoc.add_picture(str(Path.home() / Path('D:','word','malik.jpeg')),width =docx.shared.Inches(5) ,height=docx.shared.Inches(7))

mydoc.save(Path.home() / Path('D:','word','academy_3.docx'))  # if file doesnot exist will create 